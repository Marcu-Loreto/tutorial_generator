"""
Tutorial exporters — Markdown, DOCX, PDF.

Public API
----------
export_markdown(title, content_md)  → (bytes, filename)
export_docx(title, content_md)      → (bytes, filename)
export_pdf(title, content_md)       → (bytes, filename)
sanitize_filename(title)            → str

All functions:
  • Save a copy to /exports on disk.
  • Return (raw_bytes, filename) for Streamlit st.download_button.
  • Raise ExportError on any known failure.
"""

from __future__ import annotations

import io
import logging
import os
import re
import unicodedata
from datetime import datetime, timezone
from pathlib import Path

logger = logging.getLogger(__name__)

_BASE_DIR = Path(__file__).resolve().parent.parent
EXPORTS_DIR = _BASE_DIR / "exports"


class ExportError(RuntimeError):
    """Raised when an export fails for a known reason."""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _ensure_exports_dir() -> None:
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)


def sanitize_filename(title: str) -> str:
    """
    Convert a tutorial title into a safe ASCII filename stem.

    Steps: unicode-normalise → strip non-ASCII → replace unsafe chars
    with underscores → collapse repeats → truncate to 80 chars → add timestamp.
    """
    stem = unicodedata.normalize("NFKD", title)
    stem = stem.encode("ascii", "ignore").decode("ascii")
    stem = re.sub(r"[^\w\s\-]", "", stem)
    stem = re.sub(r"[\s\-]+", "_", stem).strip("_")
    stem = re.sub(r"_+", "_", stem)
    stem = stem[:80] or "tutorial"
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    return f"{stem}_{ts}"


def _save_bytes(filename: str, data: bytes) -> Path:
    """Persist bytes to EXPORTS_DIR and return the path."""
    _ensure_exports_dir()
    path = EXPORTS_DIR / filename
    path.write_bytes(data)
    logger.info("Exported: %s (%d bytes)", path, len(data))
    return path


# ---------------------------------------------------------------------------
# Markdown parser — shared by DOCX and PDF engines
# ---------------------------------------------------------------------------

def _strip_inline(text: str) -> str:
    """Remove **bold**, *italic*, `code` markers."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def _strip_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return [_strip_inline(c) for c in cells]


def _is_separator_row(line: str) -> bool:
    inner = line.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
    return inner == "" and "|" in line and "-" in line


class _Block:
    __slots__ = ("kind", "text", "level", "cells")

    def __init__(self, kind: str, text: str = "", level: int = 0,
                 cells: list[str] | None = None):
        self.kind = kind    # heading|para|bullet|ordered|code|quote|rule|table_header|table_row|blank
        self.text = text
        self.level = level  # 1-3 for headings
        self.cells = cells or []


def parse_markdown(md: str) -> list[_Block]:
    """
    Parse Markdown into a flat list of _Block objects.

    Handles: H1-H3, unordered lists, ordered lists, fenced code blocks,
    blockquotes, horizontal rules, pipe tables, normal paragraphs.
    """
    blocks: list[_Block] = []
    in_code = False
    code_lines: list[str] = []
    table_open = False

    for raw in md.splitlines():
        # ── Fenced code block ──────────────────────────────────────────
        if raw.strip().startswith("```"):
            if in_code:
                blocks.append(_Block("code", "\n".join(code_lines)))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(raw)
            continue

        s = raw.strip()

        if not s:
            blocks.append(_Block("blank"))
            table_open = False
            continue

        # ── Heading ───────────────────────────────────────────────────
        m = re.match(r"^(#{1,3})\s+(.+)$", s)
        if m:
            blocks.append(_Block("heading", _strip_inline(m.group(2)), level=len(m.group(1))))
            table_open = False
            continue

        # ── Horizontal rule ───────────────────────────────────────────
        if re.match(r"^[-*_]{3,}$", s):
            blocks.append(_Block("rule"))
            table_open = False
            continue

        # ── Table ─────────────────────────────────────────────────────
        if s.startswith("|") and s.endswith("|"):
            if _is_separator_row(s):
                continue
            cells = _strip_table_row(s)
            if not table_open:
                blocks.append(_Block("table_header", cells=cells))
                table_open = True
            else:
                blocks.append(_Block("table_row", cells=cells))
            continue
        else:
            table_open = False

        # ── Unordered list ────────────────────────────────────────────
        m = re.match(r"^[-*+]\s+(.+)$", s)
        if m:
            blocks.append(_Block("bullet", _strip_inline(m.group(1))))
            continue

        # ── Ordered list ──────────────────────────────────────────────
        m = re.match(r"^\d+\.\s+(.+)$", s)
        if m:
            blocks.append(_Block("ordered", _strip_inline(m.group(1))))
            continue

        # ── Blockquote ────────────────────────────────────────────────
        if s.startswith(">"):
            blocks.append(_Block("quote", _strip_inline(s.lstrip(">").strip())))
            continue

        # ── Normal paragraph ──────────────────────────────────────────
        blocks.append(_Block("para", _strip_inline(s)))

    if in_code and code_lines:
        blocks.append(_Block("code", "\n".join(code_lines)))

    return blocks


# ---------------------------------------------------------------------------
# 1. Markdown export
# ---------------------------------------------------------------------------

def export_markdown(title: str, content_md: str) -> tuple[bytes, str]:
    """
    Save the tutorial as a UTF-8 .md file.

    Returns:
        (bytes, filename)

    Raises:
        ExportError
    """
    if not content_md or not content_md.strip():
        raise ExportError("O conteúdo está vazio — impossível exportar.")

    stem = sanitize_filename(title)
    filename = f"{stem}.md"
    data = content_md.encode("utf-8")

    try:
        _save_bytes(filename, data)
    except OSError as exc:
        raise ExportError(f"Falha ao salvar Markdown: {exc}") from exc

    return data, filename


# ---------------------------------------------------------------------------
# 2. DOCX export (python-docx)
# ---------------------------------------------------------------------------

def export_docx(title: str, content_md: str) -> tuple[bytes, str]:
    """
    Generate a Word (.docx) document from Markdown.

    Returns:
        (bytes, filename)

    Raises:
        ExportError, ImportError
    """
    if not content_md or not content_md.strip():
        raise ExportError("O conteúdo está vazio — impossível exportar.")

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise ImportError("python-docx não instalado. Execute: pip install python-docx") from exc

    doc = Document()

    # Metadata
    doc.core_properties.title = title
    doc.core_properties.author = "TutorialGen"

    # Page margins
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(3.0)

    # ── Helpers ────────────────────────────────────────────────────────────

    def _add_code_para(line: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line or " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        pPr.append(shd)

    def _add_rule() -> None:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "AAAAAA")
        pBdr.append(bot)
        pPr.append(pBdr)

    def _add_table(headers: list[str], rows: list[list[str]]) -> None:
        if not headers:
            return
        cols = len(headers)
        tbl = doc.add_table(rows=1 + len(rows), cols=cols)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            runs = hdr[i].paragraphs[0].runs
            if runs:
                runs[0].bold = True
        for ri, row_data in enumerate(rows):
            dcells = tbl.rows[ri + 1].cells
            for ci, val in enumerate(row_data):
                if ci < cols:
                    dcells[ci].text = val
        doc.add_paragraph()

    # ── Render ─────────────────────────────────────────────────────────────
    blocks = parse_markdown(content_md)
    tbl_headers: list[str] = []
    tbl_rows: list[list[str]] = []
    ordered_idx = 0

    def flush_tbl():
        nonlocal tbl_headers, tbl_rows
        if tbl_headers:
            _add_table(tbl_headers, tbl_rows)
        tbl_headers = []
        tbl_rows = []

    for block in blocks:
        if block.kind not in ("table_header", "table_row"):
            flush_tbl()

        if block.kind == "heading":
            ordered_idx = 0
            doc.add_heading(block.text, level=min(block.level, 9))

        elif block.kind == "para":
            ordered_idx = 0
            doc.add_paragraph(block.text)

        elif block.kind == "bullet":
            doc.add_paragraph(block.text, style="List Bullet")

        elif block.kind == "ordered":
            ordered_idx += 1
            doc.add_paragraph(f"{ordered_idx}. {block.text}")

        elif block.kind == "code":
            for line in (block.text.splitlines() or [""]):
                _add_code_para(line)

        elif block.kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(block.text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        elif block.kind == "rule":
            ordered_idx = 0
            _add_rule()

        elif block.kind == "table_header":
            tbl_headers = block.cells
            tbl_rows = []

        elif block.kind == "table_row":
            tbl_rows.append(block.cells)

        # blank → skip

    flush_tbl()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    data = buf.read()

    stem = sanitize_filename(title)
    filename = f"{stem}.docx"
    try:
        _save_bytes(filename, data)
    except OSError as exc:
        raise ExportError(f"Falha ao salvar DOCX: {exc}") from exc

    return data, filename


# ---------------------------------------------------------------------------
# 3. PDF export (ReportLab Platypus)
# ---------------------------------------------------------------------------

def _escape_pdf(text: str) -> str:
    """Escape characters that break ReportLab's XML parser."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def export_pdf(title: str, content_md: str) -> tuple[bytes, str]:
    """
    Generate a PDF from Markdown using ReportLab Platypus.

    Returns:
        (bytes, filename)

    Raises:
        ExportError, ImportError
    """
    if not content_md or not content_md.strip():
        raise ExportError("O conteúdo está vazio — impossível exportar.")

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            HRFlowable, Preformatted, Table, TableStyle,
            ListFlowable, ListItem,
        )
    except ImportError as exc:
        raise ImportError("reportlab não instalado. Execute: pip install reportlab") from exc

    buf = io.BytesIO()
    page_w, page_h = A4
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=3 * cm, rightMargin=3 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=title, author="TutorialGen",
    )

    base = getSampleStyleSheet()
    S = {
        "h1": ParagraphStyle("TGH1", parent=base["Heading1"],
                              fontSize=20, spaceAfter=12,
                              textColor=colors.HexColor("#1a1a2e")),
        "h2": ParagraphStyle("TGH2", parent=base["Heading2"],
                              fontSize=15, spaceAfter=8,
                              textColor=colors.HexColor("#16213e")),
        "h3": ParagraphStyle("TGH3", parent=base["Heading3"],
                              fontSize=12, spaceAfter=6,
                              textColor=colors.HexColor("#0f3460")),
        "normal": ParagraphStyle("TGNormal", parent=base["Normal"],
                                 fontSize=10, leading=15, spaceAfter=6),
        "bullet": ParagraphStyle("TGBullet", parent=base["Normal"],
                                 fontSize=10, leading=14,
                                 leftIndent=18, spaceAfter=3),
        "ordered": ParagraphStyle("TGOrdered", parent=base["Normal"],
                                  fontSize=10, leading=14,
                                  leftIndent=18, spaceAfter=3),
        "code": ParagraphStyle("TGCode", parent=base["Code"],
                               fontSize=8.5, leading=12, fontName="Courier",
                               backColor=colors.HexColor("#F5F5F5"),
                               leftIndent=12, rightIndent=12,
                               spaceBefore=4, spaceAfter=4,
                               textColor=colors.HexColor("#222222")),
        "quote": ParagraphStyle("TGQuote", parent=base["Normal"],
                                fontSize=10, leading=14,
                                leftIndent=20, spaceAfter=6,
                                textColor=colors.HexColor("#555555")),
    }

    avail_width = page_w - 6 * cm
    story: list = []

    # Table & list accumulators
    tbl_headers: list[str] = []
    tbl_rows_pdf: list[list[str]] = []
    bullet_buf: list[str] = []
    ordered_buf: list[str] = []

    def flush_bullets():
        nonlocal bullet_buf
        if not bullet_buf:
            return
        items = [
            ListItem(Paragraph(_escape_pdf(t), S["bullet"]),
                     leftIndent=18,
                     bulletColor=colors.HexColor("#0f3460"))
            for t in bullet_buf
        ]
        story.append(ListFlowable(items, bulletType="bullet", start="•"))
        story.append(Spacer(1, 4))
        bullet_buf = []

    def flush_ordered():
        nonlocal ordered_buf
        if not ordered_buf:
            return
        items = [
            ListItem(Paragraph(_escape_pdf(t), S["ordered"]), leftIndent=18)
            for t in ordered_buf
        ]
        story.append(ListFlowable(items, bulletType="1"))
        story.append(Spacer(1, 4))
        ordered_buf = []

    def flush_tbl_pdf():
        nonlocal tbl_headers, tbl_rows_pdf
        if not tbl_headers:
            return
        all_rows = [tbl_headers] + tbl_rows_pdf
        n_cols = len(tbl_headers)
        col_w = avail_width / max(n_cols, 1)
        data = [
            [Paragraph(_escape_pdf(cell), S["normal"]) for cell in row]
            for row in all_rows
        ]
        tbl = Table(data, colWidths=[col_w] * n_cols, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND",   (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
            ("TEXTCOLOR",    (0, 0), (-1, 0), colors.white),
            ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, 0), 9),
            ("ALIGN",        (0, 0), (-1, -1), "LEFT"),
            ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#F8F8F8")]),
            ("GRID",         (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("TOPPADDING",   (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
            ("LEFTPADDING",  (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 8))
        tbl_headers = []
        tbl_rows_pdf = []

    blocks = parse_markdown(content_md)

    for block in blocks:
        is_tbl = block.kind in ("table_header", "table_row")
        is_lst = block.kind in ("bullet", "ordered")

        if not is_tbl:
            flush_tbl_pdf()
        if not is_lst:
            flush_bullets()
            flush_ordered()

        if block.kind == "heading":
            key = f"h{min(block.level, 3)}"
            story.append(Paragraph(_escape_pdf(block.text), S[key]))

        elif block.kind == "para":
            story.append(Paragraph(_escape_pdf(block.text), S["normal"]))

        elif block.kind == "bullet":
            bullet_buf.append(block.text)

        elif block.kind == "ordered":
            ordered_buf.append(block.text)

        elif block.kind == "code":
            for line in (block.text.splitlines() or [""]):
                story.append(Preformatted(line or " ", S["code"]))

        elif block.kind == "quote":
            story.append(Paragraph(
                f"<i>{_escape_pdf(block.text)}</i>", S["quote"]
            ))

        elif block.kind == "rule":
            story.append(Spacer(1, 4))
            story.append(HRFlowable(
                width="100%", thickness=0.5,
                color=colors.HexColor("#AAAAAA"),
            ))
            story.append(Spacer(1, 4))

        elif block.kind == "table_header":
            tbl_headers = block.cells
            tbl_rows_pdf = []

        elif block.kind == "table_row":
            tbl_rows_pdf.append(block.cells)

        elif block.kind == "blank":
            story.append(Spacer(1, 6))

    # Final flush
    flush_bullets()
    flush_ordered()
    flush_tbl_pdf()

    try:
        doc.build(story)
    except Exception as exc:
        raise ExportError(f"Falha ao gerar PDF: {exc}") from exc

    buf.seek(0)
    data = buf.read()

    stem = sanitize_filename(title)
    filename = f"{stem}.pdf"
    try:
        _save_bytes(filename, data)
    except OSError as exc:
        raise ExportError(f"Falha ao salvar PDF: {exc}") from exc

    return data, filename
